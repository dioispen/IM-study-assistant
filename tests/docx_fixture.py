"""Builds .docx fixtures at test time instead of committing them.

docs/corpus-sources.md asks of a committed fixture that it be deterministic and
reviewable in a diff. A .docx is a zip of XML carrying timestamps, so it is
neither: a reviewer reading the diff of one sees a wall of base64 and has to
take the test's word for what is in the file. Written from the block list
below, the fixture's headings and prose sit in the test that depends on them,
where the diff shows them.

What this cannot cover is the shape of a .docx some other word processor
writes, since it is python-docx reading back what python-docx wrote. The reader
is deliberately narrow about what it trusts in return -- built-in heading
styles and paragraph text, nothing that depends on how a particular writer lays
the package out.
"""

import datetime
import zipfile
from pathlib import Path

from docx import Document

# A block is (style, content). `TABLE` takes a list of rows of cell strings;
# every other style takes one paragraph of text, `BODY` being unstyled prose
# and "Heading 1".."Heading 9" the built-in heading styles Word applies.
BODY = None
TABLE = "table"


def write_docx(path: Path, blocks: list[tuple[str | None, object]]) -> Path:
    document = Document()

    for style, content in blocks:
        if style == TABLE:
            rows = list(content)
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            for row, cells in zip(table.rows, rows, strict=True):
                for cell, text in zip(row.cells, cells, strict=True):
                    cell.text = text
        else:
            document.add_paragraph(str(content), style=style)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def resave_untouched(source: Path, target: Path) -> Path:
    """`source` opened and saved again with nothing edited but the save stamp.

    What opening a note in Word and closing it does. python-docx writes a
    deterministic package -- same content, same bytes -- so re-saving through
    it alone would prove nothing; stamping `dcterms:modified` is the part of a
    real save that moves the bytes while leaving every word where it was.
    """
    document = Document(str(source))
    document.core_properties.modified = datetime.datetime(2026, 8, 21, 9, 30)
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    return target


def corrupt_a_part(source: Path, target: Path, part: str = "word/document.xml") -> Path:
    """`source` repackaged with one part replaced by text that is not XML.

    The failure a truncated sync or a half-written save actually produces, and
    the one no reader sees coming: the package opens, the part is there, and it
    falls over on being parsed rather than on being found. Random bytes fail
    earlier and by an easier route.
    """
    with zipfile.ZipFile(source) as original, zipfile.ZipFile(target, "w") as rebuilt:
        for item in original.infolist():
            data = original.read(item.filename)
            rebuilt.writestr(item, b"<w:document><not closed" if item.filename == part else data)
    return target
