"""Recovering a Document's text and its structure from the file it arrived as
(ADR-0006, ADR-0007).

A student's notes are not one file format, and which format something happens
to be saved in says nothing about what is in it. So ingestion routes on the
format and nothing else: a mixed folder ingests in one run with no flag naming
which file is which (PLAN.md §第 3 週).

Routing here rather than on Source type deliberately. Source type answers
"whose words am I reading?" (CONTEXT.md) -- a note written by the student and a
textbook chapter are different Source types whether or not they are the same
format, and a note is a note whether it was saved as Markdown or as Word. Two
different questions, and folding them together would make `--source-type note`
a claim about a file extension.

What a format decides is which *structure* it can offer, and that is what
picks the chunking path (ADR-0007). `.md` and `.docx` have headings, so they
yield the Sections `core/chunking.py` states the structured rules over. A
`.pdf` has pages and no headings a reader can trust, so it yields Extents for
the windowed path -- a paper, and a PDF note with it, because the split is by
chunking path and not by the word "note" (PLAN.md §5.1).

What a format does not get to decide is the chunking rules themselves. Every
reader here yields one of those two shapes and stops; merging, splitting and
windowing live in one place each, so a docx note cannot chunk by rules that
have drifted from the ones a Markdown note chunks by, and a paper cannot be
windowed by a size of its reader's own choosing.
"""

import re
from dataclasses import dataclass
from pathlib import Path

import docx
import pymupdf
from docx.table import Table

from core.chunking import Extent, Section, parse_markdown_sections


class ExtractionError(Exception):
    """A Document's text could not be recovered in a usable form.

    Raised per Document and caught by the run (ingestion/common.py), never
    propagated out of `ingest_folder` -- one unreadable file must not cost the
    corpus every file after it in the walk.
    """


@dataclass(frozen=True)
class StructuredDocument:
    """A Document whose format told the reader where its author changed subject.

    Sections for `chunk_sections`, and the only shape the structured rules can
    be stated over: they are about headed prose, so a Document with no headings
    cannot be handed to them and be chunked by anything but accident.

    `text` is what content_hash is taken over, and it is per-format on purpose
    (ADR-0006). Markdown's is the file's own text, because that is what every
    Markdown Document already in a registry was hashed as and a different
    rendering here would re-embed the whole corpus for a change nobody made
    (ADR-0001). A docx's is the extracted rendering instead, because its bytes
    are a zip: opening a note in Word and saving it untouched rewrites
    timestamps and can reorder parts, so hashing the package would re-ingest a
    note whose words never moved. Hashing what was extracted asks the question
    the skip actually wants answered -- have the note's words changed? --
    rather than a question about the file.
    """

    text: str
    sections: list[Section]


@dataclass(frozen=True)
class UnstructuredDocument:
    """A Document whose format offers a place to cite but not a place to split.

    Extents for `chunk_windows`. A PDF's pages are the case this exists for:
    page 7 is a perfectly good answer to "where in the Document did this come
    from?" and no answer at all to "where does the author change subject?", so
    it names Chunks without bounding them.

    `text` is the extracted rendering, for the reason a docx's is: a PDF's
    bytes carry a creation timestamp and an object layout, so the same paper
    saved twice is not the same file, and hashing it would re-embed a paper
    whose words never moved.
    """

    text: str
    extents: list[Extent]


# What a reader hands back: the two structures a format can offer, and the
# choice of chunking path with it. A union rather than one type with both
# fields, so that "a Document is structured or it is not" is a question the
# type answers -- ingestion/common.py matches on it -- rather than an invariant
# every caller has to remember to check.
ExtractedDocument = StructuredDocument | UnstructuredDocument


# Word's built-in heading styles, by name ("Heading 2") or by style id
# ("Heading2"). Both are checked because the name is what a document written in
# a localised Word carries in its style definitions and the id is what the
# built-in style is keyed by; either identifies the same style.
_HEADING_STYLE_RE = re.compile(r"^heading\s*([1-9])$", re.IGNORECASE)

# Markdown's `#` runs out at six, and Locators from the two formats have to
# mean the same thing, so a Word Heading 7 nests at the same depth a Markdown
# note's deepest heading does rather than opening a level Markdown has no way
# to express.
_MAX_HEADING_LEVEL = 6

# The cells of one table row, joined as they are read. A table in a study note
# is rows of related terms, and running them together with no separator would
# glue the last word of one cell to the first word of the next -- which the CJK
# token measure then reads as a single word that occurs nowhere in the note.
_CELL_SEPARATOR = " | "


def _heading_level(paragraph) -> int | None:
    """The heading level `paragraph` is styled at, or None if it is body text."""
    style = paragraph.style
    if style is None:
        return None
    for identifier in (style.name, style.style_id):
        match = _HEADING_STYLE_RE.match(identifier or "")
        if match:
            return min(int(match.group(1)), _MAX_HEADING_LEVEL)
    return None


def _table_lines(table: Table) -> list[str]:
    """A table's rows as body lines, one line per row.

    Rendered rather than dropped: `Document.paragraphs` does not reach inside a
    table, so a reader that walks only paragraphs loses a comparison table
    whole -- and loses it silently, with the run reporting an ordinary success
    and the note answering nothing it was kept for.
    """
    return [line for row in table.rows if (line := _row_text(row))]


def _row_text(row) -> str:
    """One table row as a line, each cell appearing once.

    `row.cells` yields a cell per grid column, so a cell merged across three
    columns comes back three times -- a header reading `比較表` renders as
    `比較表 | 比較表 | 比較表`, in the Chunk a reader is shown and in the hash.
    Merged headers are ordinary in exactly the comparison tables this exists to
    keep, so the repeats are dropped by identity of the underlying cell: two
    entries that are the same cell are the same cell, whatever it holds.

    `_tc` because python-docx offers nothing public that answers it. A cell
    merged vertically still appears in each of its rows, which is wanted -- a
    row is read as a row, and its label belongs on every line it labels.
    """
    cells = []
    for cell in row.cells:
        if not cells or cell._tc is not cells[-1]._tc:
            cells.append(cell)
    return _CELL_SEPARATOR.join(cell.text.strip() for cell in cells).strip()


def _docx_sections(path: Path) -> list[Section]:
    """The Sections of a Word document, in document order.

    `iter_inner_content` rather than `.paragraphs`, so a table is seen where it
    sits and lands under the heading it was written beneath rather than at the
    end of the note or nowhere at all.
    """
    try:
        document = docx.Document(str(path))
        blocks = list(document.iter_inner_content())
    except Exception as error:
        # Deliberately every exception, not a tuple of the ones a .docx is
        # expected to fail with. A malformed package reaches here through
        # whichever library is unwrapping it at the time -- OpcError for a file
        # that is not a Word package, BadZipFile for a truncated sync,
        # XMLSyntaxError for a part that opens and then does not parse -- and
        # the list is a guess about a third-party parser's internals. Guessing
        # short is not one bad note reported: it is `ingest_folder` aborting
        # mid-walk, so every note after this one in the corpus goes unread,
        # retirement never runs, and the CLI prints a traceback in place of a
        # report. One unreadable file must cost the run only itself.
        raise ExtractionError(f"is not a readable .docx ({error})") from error

    sections: list[tuple[tuple[str, ...], list[str]]] = [((), [])]
    stack: list[str] = []

    for block in blocks:
        if isinstance(block, Table):
            sections[-1][1].extend(_table_lines(block))
            continue

        level = _heading_level(block)
        heading = block.text.strip()
        if level is not None and heading:
            # The same stack Markdown's `#` levels drive: a heading replaces
            # everything at its level and below, so a Heading 1 after a
            # Heading 2 opens a new path instead of nesting under the old one.
            stack[level - 1 :] = [heading]
            sections.append((tuple(stack), []))
        elif block.text.strip():
            sections[-1][1].append(block.text)

    # A heading with nothing under it is a place in the Document rather than a
    # span of it, and has no text to embed -- dropped here exactly as
    # parse_markdown_sections drops it, so the two formats put the same shape
    # of note into the registry.
    return [
        Section(heading_path=path_, body=body)
        for path_, lines in sections
        if (body := "\n".join(lines).strip())
    ]


def _as_hashable_text(sections: list[Section]) -> str:
    """Sections rendered as one string, for hashing and for nothing else.

    Never parsed back and never chunked -- the Sections themselves are chunked.
    That is what lets a docx body line reading `#include <stdio.h>` stay prose:
    it passes through this rendering, but nothing reads the rendering back as
    Markdown afterwards.

    Each Section is written under its whole Locator rather than under the one
    heading that opened it. An outer heading with no prose directly beneath it
    contributes no Section of its own, so rendering only the innermost heading
    would leave that outer heading out of the hash entirely -- and renaming it
    would move every Locator in the note while content_hash stayed where it
    was, which is exactly the stale generation no later run can detect that
    ADR-0001 exists to prevent.
    """
    return "\n\n".join(
        f"{'#' * len(section.heading_path)} {section.locator}\n\n{section.body}".strip()
        for section in sections
    )


# How a page is cited. "p. 7" rather than a bare number, because a Locator is
# written for a reader who wants to go look (CONTEXT.md) and a lone "7" beside
# a heading path like "死結 › 偵測" reads as neither a page nor a section.
def _page_locator(number: int) -> str:
    return f"p. {number}"


def _pdf_extents(path: Path) -> list[Extent]:
    """The pages of a PDF, in order, one Extent each.

    `get_text()` in the order the page lays its text out, with no sorting and
    no attempt to find headings. A paper arrives two-column, and both of the
    things a reader might do about that make it worse: sorting blocks by
    position interleaves the two columns line by line, and inferring headings
    from font size invents a Locator citing a place in the paper that is not
    there. The page number is a citation that holds whatever the column order
    turned out to be, which is why this path cites pages (ADR-0007). Extraction
    quality on a real paper is still checked by hand before it is trusted
    (PLAN.md §5.3).

    Empty pages are kept rather than skipped, so a page number is the number
    printed on the page. Renumbering past a full-page figure would shift every
    Locator after it and cite the wrong page for the rest of the paper; the
    empty Extent is dropped later by the chunker, which cites nothing it has
    no text for.
    """
    try:
        with pymupdf.open(str(path)) as document:
            if document.needs_pass:
                # Named rather than left to come back as a Document with no
                # text: the two failures ask opposite things of the corpus
                # owner. One wants OCR; this one wants the password they
                # already have.
                raise ExtractionError(
                    "is password-protected, so its text cannot be read "
                    "(save an unprotected copy into the corpus)"
                )
            extents = [
                Extent(locator=_page_locator(number), text=page.get_text())
                for number, page in enumerate(document, start=1)
            ]
    except ExtractionError:
        raise
    except Exception as error:
        # Every exception, for the reason _docx_sections gives at length: what
        # a malformed PDF raises is a fact about PyMuPDF's internals -- a
        # FileDataError for a truncated sync, something else for a file that
        # opens and then does not parse -- and guessing the list short costs
        # the run every file after this one rather than costing it this one.
        raise ExtractionError(f"is not a readable .pdf ({error})") from error

    if not any(extent.text.strip() for extent in extents):
        # A scanned paper: every page an image and not a word between them.
        raise ExtractionError(
            "has no text layer to extract (a scan needs OCR before it can be "
            "ingested)"
        )
    return extents


def _as_hashable_extents(extents: list[Extent]) -> str:
    """Extents rendered as one string, for hashing and for nothing else.

    Each page is written under its own Locator, so that the same words spread
    across a different number of pages hash differently. They have to: every
    Locator after the break has moved, and a hash blind to that leaves the
    corpus citing page numbers the file no longer has, with no later run able
    to detect it (ADR-0001).
    """
    return "\n\n".join(
        f"[{extent.locator}]\n{extent.text.strip()}" for extent in extents
    )


def _extract_pdf(path: Path) -> UnstructuredDocument:
    extents = _pdf_extents(path)
    return UnstructuredDocument(
        text=_as_hashable_extents(extents), extents=extents
    )


def _extract_markdown(path: Path) -> StructuredDocument:
    # read_text, not read_bytes().decode(): it translates line endings, so a
    # note's content_hash survives a CRLF-vs-LF checkout instead of the whole
    # corpus re-embedding once on a different machine.
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as error:
        raise ExtractionError(f"not valid utf-8 ({error.reason})") from error
    return StructuredDocument(text=raw, sections=parse_markdown_sections(raw))


def _extract_docx(path: Path) -> StructuredDocument:
    sections = _docx_sections(path)
    return StructuredDocument(text=_as_hashable_text(sections), sections=sections)


_EXTRACTORS = {
    ".md": _extract_markdown,
    ".docx": _extract_docx,
    ".pdf": _extract_pdf,
}

# What the walk looks for: every format with a reader here, whichever shape
# that reader hands back. Lower-cased, and matched against a lower-cased
# suffix, so a file saved as `NOTES.DOCX` is the same Document on Windows and
# on Linux rather than one that exists in one machine's corpus and not the
# other's.
DOCUMENT_SUFFIXES = frozenset(_EXTRACTORS)


def extract_document(path: Path) -> ExtractedDocument:
    """One Document's text and structure, routed by the format its file is in.

    Which of the two shapes comes back is the format's answer to whether it
    knows where the author changed subject, and it is what picks the chunking
    path downstream.

    Raises ExtractionError, never anything a caller has to know a reader's
    library to catch.
    """
    extractor = _EXTRACTORS.get(path.suffix.lower())
    if extractor is None:
        # The walk filters on DOCUMENT_SUFFIXES, so reaching this is a bug
        # rather than a bad Document -- still ExtractionError, because the
        # alternative is a traceback that costs the folder every file after
        # this one.
        raise ExtractionError(
            f"has no reader for {path.suffix or 'a file with no suffix'!r}"
        )
    try:
        return extractor(path)
    except OSError as error:
        # Locked by an editor, mid-sync, deleted since the walk -- per-file
        # accidents, and no reason to cost the folder every note after this one.
        raise ExtractionError(f"could not be read ({error.strerror or error})") from error
